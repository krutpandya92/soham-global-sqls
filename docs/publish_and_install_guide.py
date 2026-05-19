"""
Generates publish_and_install_guide.docx — a step-by-step, fully-documented
guide for publishing global_sqls to npm and installing it as an MCP server
at user scope, project scope, and local scope. Every command is explained
with all flags and arguments.

Reuses the visual style of npm_link_and_mcp_command_resolution.py.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def shade_cell(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    run = p.runs[0] if p.runs else p.add_run()
    run.text = text
    run.font.size = Pt(11)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), 'Consolas')
    rFonts.set(qn('w:hAnsi'), 'Consolas')
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), 'F4F4F4')
    pPr.append(shd)
    return p


def add_table(doc, headers, rows, header_fill='1F3A5F'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        shade_cell(hdr[i], header_fill)
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
    return table


def main():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ---------------- TITLE ----------------
    title = doc.add_heading('Publishing & Installing global_sqls as an MCP Server', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run('Step-by-step: npm publish, user-scope, project-scope, local-scope install')
    r.italic = True
    r.font.size = Pt(12)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run('Working example: global_sqls   |   Platform: Windows + PowerShell\n').font.size = Pt(10)
    meta.add_run('Author: krutpandya92   |   Date: 2026-05-20').font.size = Pt(10)

    doc.add_paragraph()

    # ---------------- 1. OVERVIEW ----------------
    add_heading(doc, '1. Overview', 1)
    add_para(doc,
        "This guide covers the full life-cycle of shipping the global_sqls MCP server "
        "so other people can install and run it. There are two distribution paths and "
        "three install scopes; pick the combination that matches your audience.")

    add_heading(doc, '1.1 Distribution paths', 2)
    add_table(doc,
        ['Path', 'When to use', 'Who can install'],
        [
            ['npm registry (public)', 'You want the world to `npx -y global-sqls`.', 'Anyone with Node.js'],
            ['npm scoped package (@you/global-sqls)', 'Public, but namespaced under your account.', 'Anyone with Node.js'],
            ['Private npm / GitHub Packages', 'Internal-only, paid org or company registry.', 'Members of your org with an auth token'],
            ['Remote HTTP/SSE server', 'You centralize DB credentials and expose tools over network.', 'Anyone you give the URL + token to'],
        ])

    add_heading(doc, '1.2 Install scopes (Claude Code)', 2)
    add_table(doc,
        ['Scope', 'Config file', 'Visibility'],
        [
            ['user', '~/.claude.json (Windows: %USERPROFILE%\\.claude.json)', 'You, all projects on this machine'],
            ['project', '.mcp.json at repo root (commit it)', 'Anyone who clones the repo'],
            ['local', '.claude/settings.local.json (do NOT commit)', 'Only you, only in this project'],
        ])

    # ---------------- 2. PRE-PUBLISH PREP ----------------
    add_heading(doc, '2. Pre-publish: prepare package.json', 1)
    add_para(doc,
        "Before the first `npm publish` you must fix a handful of metadata fields. "
        "These are required for a discoverable, installable, cross-platform package.")

    add_heading(doc, '2.1 Required edits to package.json', 2)
    add_bullet(doc, 'Rename "global_sqls" → "global-sqls". npm allows underscores but the convention is lowercase + hyphens; many registries and docs assume hyphens.')
    add_bullet(doc, 'Bump "version" to "1.0.0" — the first public release.')
    add_bullet(doc, 'Add "files": ["build","README.md","LICENSE"] so only built output ships (not src/, tests/, .git/).')
    add_bullet(doc, 'Add "repository", "homepage", "bugs", "license", "keywords" so npm search and the package page work.')
    add_bullet(doc, 'Add "prepublishOnly" script so a stale or untested build can never be published.')

    add_heading(doc, '2.2 Concrete patched package.json', 2)
    add_code(doc,
        '{\n'
        '  "name": "global-sqls",\n'
        '  "version": "1.0.0",\n'
        '  "description": "Multi-engine SQL MCP server (MSSQL, MySQL, PostgreSQL, Oracle) with named connection profiles and audit logging",\n'
        '  "type": "module",\n'
        '  "main": "build/index.js",\n'
        '  "bin": { "global-sqls": "./build/index.js" },\n'
        '  "files": ["build", "README.md", "LICENSE"],\n'
        '  "scripts": {\n'
        '    "build": "tsc",\n'
        '    "prepublishOnly": "npm run clean && npm run build && npm test",\n'
        '    "clean": "node -e \\"require(\'fs\').rmSync(\'build\',{recursive:true,force:true})\\""\n'
        '  },\n'
        '  "keywords": ["mcp","sql","mssql","mysql","postgres","oracle","model-context-protocol"],\n'
        '  "repository": { "type": "git", "url": "git+https://github.com/krutpandya92/global_sqls.git" },\n'
        '  "homepage": "https://github.com/krutpandya92/global_sqls#readme",\n'
        '  "bugs": { "url": "https://github.com/krutpandya92/global_sqls/issues" },\n'
        '  "license": "MIT",\n'
        '  "engines": { "node": ">=18.0.0" }\n'
        '}')

    add_heading(doc, '2.3 Why each field matters', 2)
    add_table(doc,
        ['Field', 'Purpose'],
        [
            ['name', 'Globally unique on npm. Used in install URLs and require() calls.'],
            ['version', 'SemVer. npm refuses to publish over an existing version.'],
            ['bin', 'Maps CLI name → entry script. After install, npm creates a launcher on PATH.'],
            ['main', 'Default require/import entry when the package is consumed as a library.'],
            ['files', 'Whitelist of what to include in the tarball. Without it everything except .gitignore matches.'],
            ['prepublishOnly', 'Runs before `npm publish` only (not on install). Last guard against publishing broken code.'],
            ['engines', 'Advisory minimum Node version. With strict-engines or pnpm, it is enforced.'],
            ['repository / homepage / bugs', 'Shown on the npm package page; help users find source and report issues.'],
            ['keywords', 'Powers npm search.'],
        ])

    add_heading(doc, '2.4 Required runtime detail: shebang', 2)
    add_para(doc,
        'The first line of build/index.js MUST be the shebang below, or the bin will not run '
        'on macOS / Linux (Windows uses .cmd shims that ignore it, but POSIX systems require it):')
    add_code(doc, '#!/usr/bin/env node')
    add_para(doc,
        'Source it in src/index.ts so tsc preserves it. Also chmod the built file in your build step '
        'if you publish from a POSIX CI runner: `chmod +x build/index.js`.')

    add_heading(doc, '2.5 LICENSE file', 2)
    add_para(doc, 'Create a LICENSE file in repo root. MIT example:')
    add_code(doc,
        'MIT License\n\n'
        'Copyright (c) 2026 krutpandya92\n\n'
        'Permission is hereby granted, free of charge, to any person obtaining a copy\n'
        'of this software and associated documentation files (the "Software"), to deal\n'
        'in the Software without restriction, ...')

    # ---------------- 3. NPM ACCOUNT & LOGIN ----------------
    add_heading(doc, '3. npm account & login', 1)
    add_para(doc, 'You need an npm account before you can publish. Create one once, then log in per machine.')

    add_heading(doc, '3.1 Create the account', 2)
    add_bullet(doc, 'Sign up at https://www.npmjs.com/signup.')
    add_bullet(doc, 'Enable two-factor authentication (2FA) — recommended; npm can require a TOTP code at publish time.')

    add_heading(doc, '3.2 Log in from your shell', 2)
    add_code(doc, 'npm login')
    add_para(doc, 'Explanation of `npm login`:', bold=True)
    add_bullet(doc, 'Opens the default browser to npmjs.com, then stores an auth token in ~/.npmrc.')
    add_bullet(doc, 'Default registry is https://registry.npmjs.org/. Override with --registry=<url> for private registries.')
    add_bullet(doc, '--scope=@your-scope tells npm to authenticate for a scoped package.')
    add_bullet(doc, '--auth-type=web (default) opens browser; --auth-type=legacy prompts for username/password/OTP in the terminal.')

    add_heading(doc, '3.3 Verify you are logged in', 2)
    add_code(doc, 'npm whoami')
    add_para(doc, 'Prints your npm username if a valid token is present, otherwise exits with an error.')

    # ---------------- 4. PUBLISH ----------------
    add_heading(doc, '4. Publish to npm', 1)

    add_heading(doc, '4.1 Dry run first', 2)
    add_code(doc, 'npm publish --dry-run')
    add_para(doc, 'Flags explained:', bold=True)
    add_bullet(doc, '--dry-run: simulates publish; prints the tarball contents and total size; nothing is uploaded.')
    add_para(doc,
        'Inspect the file list — anything you do not want shipped (src/, tests/, .env, screenshots) '
        'means you need to fix "files" in package.json or add a .npmignore.')

    add_heading(doc, '4.2 Publish public (unscoped name)', 2)
    add_code(doc, 'npm publish')
    add_para(doc,
        'Defaults: registry = https://registry.npmjs.org/, access = public (for unscoped names), tag = latest.')

    add_heading(doc, '4.3 Publish public (scoped name like @krutpandya92/global-sqls)', 2)
    add_code(doc, 'npm publish --access public')
    add_para(doc, 'Why --access public:', bold=True)
    add_bullet(doc, 'Scoped packages default to access=restricted, which requires a paid npm org. Free public scoped packages must be explicitly marked public on first publish.')

    add_heading(doc, '4.4 All useful `npm publish` flags', 2)
    add_table(doc,
        ['Flag', 'Effect'],
        [
            ['--access public | restricted', 'Visibility for scoped packages. Unscoped names are always public.'],
            ['--tag <name>', 'Distribution tag. Default "latest". Use "beta", "next", "rc" for pre-releases — installers must opt-in with `npm i pkg@beta`.'],
            ['--dry-run', 'Show what would be published without uploading.'],
            ['--otp <code>', 'One-time 2FA code if your account requires it. Otherwise npm prompts interactively.'],
            ['--registry <url>', 'Publish to a non-default registry (private, GitHub Packages, Verdaccio).'],
            ['--provenance', 'Attach a signed SLSA provenance statement (CI-only, GitHub Actions / GitLab).'],
        ])

    add_heading(doc, '4.5 What happens on success', 2)
    add_bullet(doc, 'A tarball (.tgz) is uploaded to the registry.')
    add_bullet(doc, 'The package page goes live at https://www.npmjs.com/package/global-sqls within ~1 minute.')
    add_bullet(doc, 'Anyone can now run: `npx -y global-sqls` (downloads on first use, caches in ~/.npm/_npx).')

    add_heading(doc, '4.6 Post-publish smoke test', 2)
    add_code(doc, 'npx -y global-sqls --help')
    add_para(doc, 'Flag explanation:', bold=True)
    add_bullet(doc, '-y / --yes: auto-confirm the "Need to install ... Ok to proceed?" prompt.')
    add_bullet(doc, 'If your server has no --help flag (most MCP servers speak only over stdio), this will block waiting for JSON-RPC. Press Ctrl+C — the fact that it started without error is the smoke test.')

    # ---------------- 5. VERSIONING / UPDATES ----------------
    add_heading(doc, '5. Publishing updates', 1)
    add_para(doc, 'Every publish must use a NEW version number. Use `npm version` to bump + tag + commit in one shot.')
    add_code(doc,
        'npm version patch   # 1.0.0 -> 1.0.1  (bugfix)\n'
        'npm version minor   # 1.0.1 -> 1.1.0  (backward-compatible feature)\n'
        'npm version major   # 1.1.0 -> 2.0.0  (breaking change)\n'
        'npm publish')
    add_para(doc, '`npm version` does three things automatically:', bold=True)
    add_bullet(doc, 'Rewrites the "version" field in package.json (and package-lock.json).')
    add_bullet(doc, 'Creates a git commit with message `vX.Y.Z` (configurable via -m "...").')
    add_bullet(doc, 'Creates a git tag `vX.Y.Z`.')
    add_para(doc, 'Push commits and tags after publishing:')
    add_code(doc, 'git push && git push --tags')

    # ---------------- 6. .mcp.json EXAMPLES ----------------
    add_heading(doc, '6. .mcp.json — the universal MCP config block', 1)
    add_para(doc,
        'Once published, users add one of the snippets below to their MCP client config. '
        'The same JSON works for Claude Code, Claude Desktop, Cursor, Windsurf, Copilot CLI — '
        'only the file path differs.')

    add_heading(doc, '6.1 Minimal block', 2)
    add_code(doc,
        '{\n'
        '  "mcpServers": {\n'
        '    "global-sqls": {\n'
        '      "command": "npx",\n'
        '      "args": ["-y", "global-sqls"]\n'
        '    }\n'
        '  }\n'
        '}')

    add_heading(doc, '6.2 With env vars (recommended)', 2)
    add_code(doc,
        '{\n'
        '  "mcpServers": {\n'
        '    "global-sqls": {\n'
        '      "command": "npx",\n'
        '      "args": ["-y", "global-sqls"],\n'
        '      "env": {\n'
        '        "GLOBAL_SQLS_CONNECTIONS": "C:/Users/me/.global_sqls/connections.json",\n'
        '        "AUDIT_LOG_DIR": "C:/Users/me/.global_sqls/logs",\n'
        '        "AUDIT_LOG_VERBOSE": "1"\n'
        '      }\n'
        '    }\n'
        '  }\n'
        '}')

    add_heading(doc, '6.3 Field-by-field explanation', 2)
    add_table(doc,
        ['Field', 'Purpose'],
        [
            ['mcpServers', 'Top-level map. Each key is the server alias used in the client UI.'],
            ['command', 'The executable to spawn. "npx" runs from registry; "global-sqls" runs an installed global; an absolute path works too.'],
            ['args', 'Array of CLI arguments. "-y" auto-confirms npx; subsequent items are flags for the server itself.'],
            ['env', 'Env vars injected only into the spawned server process. Do not leak into your shell.'],
            ['cwd (optional)', 'Working directory for the server process. Defaults to the client\'s cwd.'],
        ])

    # ---------------- 7. INSTALL SCOPES ----------------
    add_heading(doc, '7. Installing per scope (Claude Code)', 1)

    add_heading(doc, '7.1 User scope — available everywhere on this machine', 2)
    add_code(doc, 'claude mcp add global-sqls -s user -- npx -y global-sqls')
    add_para(doc, 'Flag explanation:', bold=True)
    add_bullet(doc, '`mcp add` — subcommand that registers a new MCP server.')
    add_bullet(doc, '`global-sqls` — the alias the user will see inside Claude Code.')
    add_bullet(doc, '`-s user` / `--scope user` — writes to ~/.claude.json (the user-level config). Other values: project, local.')
    add_bullet(doc, '`--` — POSIX end-of-options. Everything after `--` is treated as the command to spawn, not as flags for `claude`.')
    add_bullet(doc, '`npx -y global-sqls` — the actual spawn command.')
    add_para(doc, 'Add env vars with --env (repeat per variable):')
    add_code(doc,
        'claude mcp add global-sqls -s user `\n'
        '  --env GLOBAL_SQLS_CONNECTIONS=C:/Users/me/.global_sqls/connections.json `\n'
        '  --env AUDIT_LOG_DIR=C:/Users/me/.global_sqls/logs `\n'
        '  -- npx -y global-sqls')
    add_para(doc, '(Backticks are PowerShell line continuations; in cmd.exe use `^`, in bash use `\\`.)')

    add_heading(doc, '7.2 Project scope — shared with your team via git', 2)
    add_code(doc, 'claude mcp add global-sqls -s project -- npx -y global-sqls')
    add_para(doc,
        'Writes a .mcp.json file at the repo root. Commit this file. Every teammate who clones '
        'the repo and opens it in Claude Code will be prompted to enable the server.')

    add_heading(doc, '7.3 Local scope — only you, only this project', 2)
    add_code(doc, 'claude mcp add global-sqls -s local -- npx -y global-sqls')
    add_para(doc,
        'Writes to .claude/settings.local.json. Already covered by the default .gitignore in '
        'Claude-Code-initialized projects. Good for personal credentials you do not want in shared config.')

    add_heading(doc, '7.4 Inspect / remove servers', 2)
    add_code(doc,
        'claude mcp list                 # show all registered servers across scopes\n'
        'claude mcp get global-sqls      # show full config for one server\n'
        'claude mcp remove global-sqls -s user   # delete from a specific scope')

    # ---------------- 8. OTHER CLIENTS ----------------
    add_heading(doc, '8. Other MCP clients — where to put the same JSON', 1)
    add_table(doc,
        ['Client', 'Config file', 'Notes'],
        [
            ['Claude Code (user)', '%USERPROFILE%\\.claude.json', 'Edit the "mcpServers" key.'],
            ['Claude Code (project)', '<repo>/.mcp.json', 'Commit to git.'],
            ['Claude Desktop', '%APPDATA%\\Claude\\claude_desktop_config.json', 'Restart the app after editing.'],
            ['Cursor', '%USERPROFILE%\\.cursor\\mcp.json', 'Reload window.'],
            ['Windsurf', '%USERPROFILE%\\.codeium\\windsurf\\mcp_config.json', 'Reload window.'],
            ['VS Code (GH Copilot)', '<repo>/.vscode/mcp.json', 'Per-workspace; commit if shared.'],
        ])

    # ---------------- 9. REMOTE HTTP DEPLOY ----------------
    add_heading(doc, '9. Optional: deploy as a remote HTTP server', 1)
    add_para(doc,
        'Stdio is per-user. If you want one shared instance that holds DB credentials centrally, '
        'switch the transport to HTTP/SSE and host it.')

    add_heading(doc, '9.1 Code change', 2)
    add_code(doc,
        "import { StreamableHTTPServerTransport } from '@modelcontextprotocol/sdk/server/streamableHttp.js';\n"
        "// instead of: new StdioServerTransport()\n"
        "const transport = new StreamableHTTPServerTransport({ port: 3333 });\n"
        "await server.connect(transport);")

    add_heading(doc, '9.2 Dockerfile sketch', 2)
    add_code(doc,
        'FROM node:20-alpine\n'
        'WORKDIR /app\n'
        'COPY package*.json ./\n'
        'RUN npm ci --omit=dev\n'
        'COPY build ./build\n'
        'EXPOSE 3333\n'
        'CMD ["node", "build/index.js"]')

    add_heading(doc, '9.3 Add to a client', 2)
    add_code(doc,
        'claude mcp add global-sqls -s user `\n'
        '  --transport http `\n'
        '  --header "Authorization: Bearer $env:GS_TOKEN" `\n'
        '  -- https://sql-mcp.example.com/mcp')
    add_bullet(doc, '`--transport http` — uses Streamable HTTP instead of stdio. Other value: `sse` (legacy).')
    add_bullet(doc, '`--header` — additional HTTP header per request. Repeat for multiple headers. Used for auth.')
    add_para(doc,
        'Security note: MCP itself does not authenticate users. Put a reverse proxy (Caddy, nginx, '
        'Cloudflare Access) in front, or implement bearer-token validation inside the server.')

    # ---------------- 10. TROUBLESHOOTING ----------------
    add_heading(doc, '10. Troubleshooting', 1)
    add_table(doc,
        ['Symptom', 'Likely cause / fix'],
        [
            ['`npm publish` -> 403 Forbidden', 'Name taken or you are not logged in. Try `npm whoami` and pick a scoped name.'],
            ['`npm publish` -> 402 Payment Required', 'Publishing a scoped package without --access public. Add the flag.'],
            ['`npm publish` -> EOTP', 'Account requires 2FA. Re-run with `--otp 123456`.'],
            ['Client says "Server failed to start"', 'Run the spawn command manually in a terminal; fix whatever error appears.'],
            ['`npx -y global-sqls` -> "command not found" after install', 'Shell PATH not refreshed. Open a new terminal, or check that npm prefix is on PATH (`npm config get prefix`).'],
            ['Server starts but no tools appear', 'Forgot to await server.connect(transport). Check stderr in the client log.'],
            ['Audit log not written', 'AUDIT_LOG_DIR not set or not writable. Check env vars actually reach the spawned process via your client log.'],
        ])

    # ---------------- 11. CHECKLIST ----------------
    add_heading(doc, '11. End-to-end checklist', 1)
    add_bullet(doc, 'package.json updated (name, version, files, repository, license, keywords, prepublishOnly).')
    add_bullet(doc, 'LICENSE file present.')
    add_bullet(doc, 'README documents env vars and shows an .mcp.json snippet.')
    add_bullet(doc, 'src/index.ts starts with `#!/usr/bin/env node`.')
    add_bullet(doc, '`npm run build` succeeds; build/index.js exists.')
    add_bullet(doc, '`npm publish --dry-run` shows only the expected files.')
    add_bullet(doc, '`npm login` done; `npm whoami` prints your username.')
    add_bullet(doc, '`npm publish` (or `npm publish --access public` for scoped) succeeds.')
    add_bullet(doc, '`npx -y global-sqls` starts the server without errors on a clean machine.')
    add_bullet(doc, 'Git tag pushed: `git push --tags`.')
    add_bullet(doc, 'Optional: GitHub Release created (`gh release create vX.Y.Z --generate-notes`).')
    add_bullet(doc, 'Optional: PR opened against modelcontextprotocol/servers to list global_sqls in the public registry.')

    out = 'F:/Angular/MCP_Servers/global_sqls/docs/publish_and_install_guide.docx'
    doc.save(out)
    print(f'Wrote {out}')


if __name__ == '__main__':
    main()
