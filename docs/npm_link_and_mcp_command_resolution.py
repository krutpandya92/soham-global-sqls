"""
Generates npm_link_and_mcp_command_resolution.docx — a well-documented
reference covering: what `npm link` does, where npm puts global packages
on nvm-for-windows, how MCP clients resolve the `command` field, and
how env vars in the MCP config are wired into the server.

Reuses the visual style of session_summary_generator.py (navy headings,
shaded code blocks, Consolas inline code, Light Grid tables).
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def shade_cell(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    run = p.runs[0] if p.runs else p.add_run()
    run.text = text
    run.font.size = Pt(11)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), 'Consolas')
    rFonts.set(qn('w:hAnsi'), 'Consolas')
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), 'F4F4F4')
    pPr.append(shd)
    return p


def add_table(doc, headers, rows, header_fill='1F3A5F'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        shade_cell(hdr[i], header_fill)
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
    return table


def main():
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ---------------- TITLE ----------------
    title = doc.add_heading('npm link & MCP `command` Resolution', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run('How a locally linked Node CLI becomes a runnable MCP server')
    r.italic = True
    r.font.size = Pt(12)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run('Working example: global_sqls   |   Platform: Windows + nvm-for-windows\n').font.size = Pt(10)
    meta.add_run('Author: krutpandya92   |   Date: 2026-05-19').font.size = Pt(10)

    doc.add_paragraph()

    # ---------------- 1. OVERVIEW ----------------
    add_heading(doc, '1. Overview', 1)
    add_para(doc,
        "This document explains the chain of mechanisms that lets an MCP client "
        "(Claude Code, Claude Desktop, Cursor, Windsurf, Copilot CLI) launch a "
        "Node-based MCP server by writing a short config block like:")
    add_code(doc,
        '"global_sqls": {\n'
        '  "command": "global-sqls",\n'
        '  "env": {\n'
        '    "GLOBAL_SQLS_CONNECTIONS": "C:/Users/me/.global_sqls/connections.json",\n'
        '    "AUDIT_LOG_DIR": "C:/Users/me/.global_sqls/logs"\n'
        '  }\n'
        '}')
    add_para(doc, 'Four pieces have to line up for that block to work:', bold=True)
    add_bullet(doc, '`package.json` declares a `bin` entry so npm knows what to expose as a CLI.')
    add_bullet(doc, '`npm link` (or `npm install -g`) places a symlink + shims in npm\'s global prefix.')
    add_bullet(doc, 'The global prefix folder is on the OS `PATH`, so a bare command name resolves.')
    add_bullet(doc, 'The MCP client spawns the resolved executable with the env vars injected.')

    # ---------------- 2. npm link ----------------
    add_heading(doc, '2. What `npm link` Actually Does', 1)
    add_para(doc, 'Two-step mechanism for using a local package as if it were installed from the registry.', italic=True)

    add_heading(doc, '2.1 Step 1 — In the package directory', 2)
    add_code(doc,
        'cd F:\\Angular\\MCP_Servers\\global_sqls\n'
        'npm link')
    add_para(doc, 'This does two things:', bold=True)
    add_bullet(doc, 'Reads package.json to find the `name` field and any `bin` entries.')
    add_bullet(doc, 'Creates a global symlink in npm\'s global node_modules pointing back to your source dir.')

    add_heading(doc, '2.2 Step 2 — In a consumer project (optional)', 2)
    add_code(doc,
        'cd some-other-project\n'
        'npm link global_sqls')
    add_para(doc,
        "Creates a local symlink inside that project's node_modules/global_sqls → the global symlink → "
        "your source dir. For a pure CLI / MCP server, step 2 is usually unnecessary — step 1 alone makes "
        "the binary available system-wide.")

    add_heading(doc, '2.3 Key properties', 2)
    add_bullet(doc, 'It is a symlink, not a copy — edits to source affect every linked consumer instantly.')
    add_bullet(doc, 'No automatic rebuild — if you ship compiled JS, you still need `npm run build` after edits.')
    add_bullet(doc, 'Windows symlinks need Developer Mode or an elevated shell; modern npm usually handles this.')
    add_bullet(doc, '`npm unlink -g <name>` (in the package dir) removes the global symlink.')

    # ---------------- 3. global prefix ----------------
    doc.add_page_break()
    add_heading(doc, '3. Where npm Puts the Symlink', 1)
    add_para(doc,
        "The exact location depends on how Node was installed. npm has a configurable "
        "`prefix` — the folder under which it places global packages and their bin shims.")

    add_table(doc,
        ['Node install method', 'Default `prefix`', 'Global node_modules'],
        [
            ['Official Node.js MSI', 'C:\\Users\\<you>\\AppData\\Roaming\\npm', '<prefix>\\node_modules'],
            ['nvm-for-windows (nvm4w)', 'C:\\nvm4w\\nodejs', '<prefix>\\node_modules'],
            ['Volta', 'C:\\Users\\<you>\\AppData\\Local\\Volta', 'Volta-managed shims'],
            ['Linux/Mac (system Node)', '/usr/local', '/usr/local/lib/node_modules'],
            ['Linux/Mac (nvm)', '~/.nvm/versions/node/<ver>', '<prefix>/lib/node_modules'],
        ])

    add_para(doc, 'Check yours with:', bold=True)
    add_code(doc,
        'npm config get prefix        # the prefix folder itself\n'
        'npm root -g                  # the global node_modules path\n'
        'npm ls -g --depth=0          # what is currently installed/linked there')

    add_para(doc, 'On this machine (Windows + nvm-for-windows):', bold=True)
    add_code(doc,
        'C:\\nvm4w\\nodejs                     ← prefix (also on PATH)\n'
        'C:\\nvm4w\\nodejs\\node_modules        ← global packages live here\n'
        'C:\\nvm4w\\nodejs\\node_modules\\global_sqls   ← symlink → F:\\Angular\\MCP_Servers\\global_sqls')

    add_para(doc, 'Folder-name gotcha:', bold=True)
    add_para(doc,
        "The symlink folder is named after `package.json` -> `name` (here: global_sqls with an underscore), "
        "NOT after the `bin` key (here: global-sqls with a dash). Easy to miss when grepping the global folder.")

    # ---------------- 4. bin shims ----------------
    add_heading(doc, '4. How the Symlink Becomes a Runnable Command', 1)
    add_para(doc, 'When `npm link` sees a `bin` entry, it also creates three companion shims in the prefix folder:')
    add_code(doc,
        'C:\\nvm4w\\nodejs\\global-sqls          ← Unix-style shell script (for Git Bash/MSYS)\n'
        'C:\\nvm4w\\nodejs\\global-sqls.cmd      ← Windows cmd.exe wrapper\n'
        'C:\\nvm4w\\nodejs\\global-sqls.ps1      ← PowerShell wrapper')
    add_para(doc, 'Each shim ultimately runs:', italic=True)
    add_code(doc, 'node F:\\Angular\\MCP_Servers\\global_sqls\\build\\index.js  %*')
    add_para(doc,
        "The shim is what `package.json` declares in the `bin` field. Defining bin as "
        "`\"global-sqls\": \"./build/index.js\"` tells npm:")
    add_bullet(doc, 'Name the shim `global-sqls`.')
    add_bullet(doc, 'Point it at `./build/index.js` relative to the package root.')
    add_bullet(doc, 'Run it under Node automatically (npm injects the node invocation).')

    # ---------------- 5. PATH resolution ----------------
    doc.add_page_break()
    add_heading(doc, '5. How a Bare `command` Value Resolves', 1)
    add_para(doc,
        "When the MCP client config contains `\"command\": \"global-sqls\"` (no slashes, no extension), "
        "the OS searches every folder listed in the PATH environment variable, in order, and runs the "
        "first match it finds.")

    add_para(doc, 'On Windows, the lookup also tries appending common executable extensions:', italic=True)
    add_code(doc, 'PATHEXT = .COM;.EXE;.BAT;.CMD;.PS1;.VBS;...')
    add_para(doc, 'So `global-sqls` resolves to `global-sqls.cmd` because:')
    add_bullet(doc, '`C:\\nvm4w\\nodejs` is on PATH.')
    add_bullet(doc, '`.CMD` is in PATHEXT.')
    add_bullet(doc, '`C:\\nvm4w\\nodejs\\global-sqls.cmd` exists.')

    add_heading(doc, '5.1 Common misconception', 2)
    add_para(doc,
        "The `command` value does NOT have to live in `C:\\nvm4w\\nodejs` specifically. ANY folder on PATH "
        "works. The npm prefix is just one of many folders the OS will search. Other examples on this machine:")
    add_code(doc,
        'node     → C:\\nvm4w\\nodejs\\node.exe\n'
        'git      → C:\\Program Files\\Git\\cmd\\git.exe\n'
        'code     → C:\\Users\\krutsays\\AppData\\Local\\Programs\\Microsoft VS Code\\bin\\code.cmd\n'
        'python   → C:\\Python312\\python.exe')

    add_heading(doc, '5.2 Three valid forms of `command`', 2)
    add_table(doc,
        ['Form', 'Example', 'When to use'],
        [
            ['Bare name (PATH lookup)', '"command": "global-sqls"', 'After npm link / npm install -g; convenient default.'],
            ['Absolute path', '"command": "C:\\\\nvm4w\\\\nodejs\\\\global-sqls.cmd"', 'When PATH is not inherited or the client strips env.'],
            ['Node + script', '"command": "node", "args": ["...\\\\build\\\\index.js"]', 'Dev mode — no install/link needed; direct file launch.'],
        ])

    # ---------------- 6. env block ----------------
    add_heading(doc, '6. How the `env` Block Reaches the Server', 1)
    add_para(doc,
        "Every key/value pair in `env` is injected into the spawned child process's environment, on top "
        "of (or replacing) the parent's environment. Inside the server you read them with `process.env.XXX`.")

    add_para(doc, 'Example wiring for global_sqls:', bold=True)
    add_code(doc,
        '// src/config.ts (excerpt)\n'
        'connectionsPath: optional("GLOBAL_SQLS_CONNECTIONS", "./connections.json"),\n'
        'audit: { dir: optional("AUDIT_LOG_DIR", "./logs") }')
    add_para(doc, 'Matching MCP config:', italic=True)
    add_code(doc,
        '"env": {\n'
        '  "GLOBAL_SQLS_CONNECTIONS": "C:/Users/me/.global_sqls/connections.json",\n'
        '  "AUDIT_LOG_DIR": "C:/Users/me/.global_sqls/logs"\n'
        '}')

    add_para(doc, 'Notes on Windows path encoding in JSON:', bold=True)
    add_bullet(doc, 'Forward slashes are fine — Node normalizes them on Windows.')
    add_bullet(doc, 'Backslashes work too, but must be escaped: "C:\\\\Users\\\\me\\\\..."')
    add_bullet(doc, 'Use forward slashes to keep JSON readable.')

    # ---------------- 7. end-to-end ----------------
    doc.add_page_break()
    add_heading(doc, '7. End-to-End: What Happens When the Client Starts the Server', 1)
    add_table(doc,
        ['#', 'Step', 'Mechanism'],
        [
            ['1', 'MCP client reads config block for "global_sqls"', 'JSON parse of claude_desktop_config.json (or equivalent)'],
            ['2', 'Resolves "command": "global-sqls"', 'OS PATH lookup + PATHEXT extension matching'],
            ['3', 'Finds C:\\nvm4w\\nodejs\\global-sqls.cmd', 'Shim created earlier by `npm link`'],
            ['4', 'Spawns the shim with the env block applied', 'CreateProcess (Windows) with merged environment'],
            ['5', 'Shim invokes node + build\\index.js', 'npm-generated wrapper script'],
            ['6', 'Server reads process.env.GLOBAL_SQLS_CONNECTIONS etc.', 'Standard Node env access'],
            ['7', 'Server speaks MCP protocol over stdin/stdout', 'StdioServerTransport'],
            ['8', 'Client lists available tools and is ready to call them', 'tools/list MCP request'],
        ])

    # ---------------- 8. troubleshooting ----------------
    add_heading(doc, '8. Troubleshooting Checklist', 1)
    add_table(doc,
        ['Symptom', 'Likely cause', 'Fix'],
        [
            ['Client says "command not found"', 'npm prefix not on PATH for this shell/client', 'Use absolute path: C:\\nvm4w\\nodejs\\global-sqls.cmd'],
            ['Server starts but cannot find connections.json', 'GLOBAL_SQLS_CONNECTIONS env not set / wrong path', 'Set absolute path in MCP config env block'],
            ['"Cannot find module" at startup', 'build/ folder missing or stale', 'Run npm run build in the package dir'],
            ['Edits to src/ have no effect', 'Forgot to rebuild after editing TypeScript', 'npm run build (or use tsx dev mode)'],
            ['No folder in global node_modules', 'Looking in wrong prefix (AppData vs nvm4w)', 'Run `npm root -g` to find the real path'],
            ['EPERM creating symlink on Windows', 'Symlinks require admin or Developer Mode', 'Enable Developer Mode in Windows Settings'],
            ['Two copies of a dependency', 'Linked package + consumer both install the same dep', 'Use peerDependencies or `npm install <path>`'],
        ])

    # ---------------- 9. quick reference ----------------
    add_heading(doc, '9. Quick Reference', 1)
    add_para(doc, 'Verification commands:', bold=True)
    add_code(doc,
        'npm config get prefix          # where global packages live\n'
        'npm root -g                    # equals <prefix>\\node_modules\n'
        'npm ls -g --depth=0            # everything globally installed/linked\n'
        'where global-sqls              # what PATH resolves to (Windows)\n'
        'global-sqls --help             # confirm bin shim is on PATH')

    add_para(doc, 'Lifecycle commands:', bold=True)
    add_code(doc,
        'npm link                       # in package dir: create global symlink + bin shims\n'
        'npm unlink -g global_sqls      # remove the global symlink\n'
        'npm install -g .               # alternative: install local package globally (copy)\n'
        'npm install -g global_sqls     # install published package from registry')

    add_para(doc, 'Mental model:', bold=True)
    add_bullet(doc, 'Package location = `npm root -g` (folder named after package `name`).')
    add_bullet(doc, 'Executable shims = `npm config get prefix` (folder must be on PATH).')
    add_bullet(doc, '`command` in MCP config = anything the OS can find on PATH or as an absolute path.')
    add_bullet(doc, '`env` in MCP config = process.env values inside the spawned server.')

    # ---------------- SAVE ----------------
    out = 'F:\\Angular\\MCP_Servers\\global_sqls\\docs\\npm_link_and_mcp_command_resolution.docx'
    doc.save(out)
    print(f'Wrote: {out}')


if __name__ == '__main__':
    main()
